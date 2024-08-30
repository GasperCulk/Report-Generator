from docxtpl import DocxTemplate,InlineImage
from docx.shared import Mm
from datetime import datetime
from io import BytesIO
from PIL import Image
import fitz
import os
import sys
import re
from time import time
from shutil import copyfile
from spire.doc import Document, FileFormat  # noqa: F403
from docx import Document as DocxDocument

def remove_first_paragraph(docx_path):
    doc = DocxDocument(docx_path)
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
    doc.save(docx_path)

def convert_and_modify_rtf_files(input_dir, output_dir):
    for filename in os.listdir(output_dir):
        file_path = os.path.join(output_dir, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)
        
        # Check for .rtf files and convert them to .docx
        if filename.endswith('.rtf') or filename.endswith('.Rtf'):
            document = Document()  # noqa: F405
            document.LoadFromFile(file_path)
            output_file_path = os.path.join(output_dir, re.sub(r'\.rtf$', '.docx',
                                                filename, flags=re.IGNORECASE))
            document.SaveToFile(output_file_path, FileFormat.Docx2019)  # noqa: F405
            document.Close()
            remove_first_paragraph(output_file_path)

        # Copy other files (like .docx) to the output directory
        else:
            output_file_path = os.path.join(output_dir, filename)
            copyfile(file_path, output_file_path)
        

def generateReport():
    print("Start generating report.")
    curTime = datetime.now().strftime("%d-%m-%Y_%H-%M")

    temp       = "modify\\Template_EMC_report.docx"
    input_dir  = "modify\\input_dir"
    output_dir = "modify\\output_dir"
    buffer_dir = "modify\\buffer_dir"
    
    print("Template file name:", temp)
    print("Make sure that all directories contain proper content.")
    output_name = "Generated_report_"+curTime+".docx"
    output_path = os.path.join(output_dir, output_name)
    tpl = DocxTemplate(temp)

    docList = [[],[],[],[]]
    docListName = [[],[],[],[]]

    catName = ["nl", "mds", "harmonic", "flicker"]
    category = {
        "nl": {},
        "mds": {},
        "harmonic": {},
        "flicker": {},
    }
    context = {
        "pdfList" : []
    }


    convert_and_modify_rtf_files(input_dir, buffer_dir)
    input_dir = buffer_dir


    for filename in os.listdir(input_dir):
        filepath = os.path.join(input_dir, filename)
        if filename.endswith('.pdf'):
            doc = fitz.open(filepath)
            page = doc.load_page(0)
            pix = page.get_pixmap()

            image_bytes = pix.tobytes(output="png") 
            image = Image.open(BytesIO(image_bytes))
            image_buffer = BytesIO()
            image.save(image_buffer, format="PNG")
            image_buffer.seek(0)

            pix.save(output_path, "jpeg")
            doc.close()
            context["pdfList"].append(InlineImage(tpl,image_buffer,width=Mm(150)))


        elif filename.endswith('.docx'):
            subdoc = tpl.new_subdoc(filepath)

            cat    = filename.split("#")[2].lower()
            subcat = filename.split("#")[1].lower()

            index = -1
            for i, c in enumerate(catName):
                if c == cat:
                    index = i
                    break
            if index == -1:
                print(f"Wrong name of file: {filename}")
                sys.exit(1)

            if subcat not in category[cat]:
                    category[cat][subcat] = len(category[cat])
                    docList[index].append([])
                    docListName[index].append([])

            docList[index][category[cat][subcat]].append(subdoc)
            docListName[index][category[cat][subcat]].append(subcat)


    print("\nTotal sub lists made: ")
    for i, cat in enumerate(docList):
        for j, subcat in enumerate(cat):
            context[f"title_{i}_{j}"] = docListName[i][j][0]
            context[f"subdocList_{i}_{j}"] = docList[i][j]
        #-------PRINT------------
        name = catName[i]
        print(name)
        for key, value in category[name].items():
            print("  ", len(docList[i][value]), key, f"(title_{i}_{value})")
    print("pdf",len(context["pdfList"]))
            
    tpl.render(context)
    tpl.save(output_path)
    print(f"\nNew document saved as: {output_path}")

startTime = time()
generateReport()
print("Generating report took:", time()-startTime)
input("Press Enter to exit.")
